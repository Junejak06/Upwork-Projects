import os,sys,re
import PyPDF2
import openai
import json
from docx import Document
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from langchain.prompts import ChatPromptTemplate
from langchain.prompts.chat import SystemMessage, HumanMessagePromptTemplate
from langchain.chat_models import ChatOpenAI
from PyPDF2 import PdfReader
import docx2txt

new_filename = "input_jd.pdf"
# Set your OpenAI GPT-3 API key here
api_key = "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"

def ms_checkfile_exist(filename):
    # Check if the file exists
    if os.path.exists(filename):
        # Copying the content of the specified file to "input_jd.pdf"
        with open(filename, 'rb') as original_file:
            content = original_file.read()
            with open(new_filename, 'wb') as new_file:
                new_file.write(content)
        print(f"{filename} has been copied to {new_filename}")
    else:
        print(f"Error: {filename} does not exist!")
        sys.exit("terminating program.")

def create_docx_L1(formatted_title, part1, part2, part3, filename_output="output.docx"):
    # Create a new Word document
    doc = Document()
    data = part1 #part 1
    #add main heading
    doc.add_heading(formatted_title, 0)
    #add heading
    doc.add_heading('Job Description', 1)
    # Add content from the dictionary to the document
    doc.add_paragraph("Start Date: " + data["Start Date"])
    doc.add_paragraph("Type: " + data["Type"])
    if not any(term in data["Type"].lower() for term in ["full time", "fulltime", "full-time"]):
      doc.add_paragraph("Estimated Duration: " + data["Estimated Duration"])
    doc.add_paragraph("Work Setting: " + data["Work Setting"])

    # part 2

    #add heading
    doc.add_heading('Job Requirements', 1)

    data = part2['Requirements']
    # list of strings
    for item in data:
        doc.add_paragraph(item)

    # part 3

    #add heading
    doc.add_heading('Job Responsibilities', 1)
    data = part3['Responsibilities']
    # list of strings
    for item in data:
        doc.add_paragraph(item)
    # Save the document
    doc.save(filename_output)
    return formatted_title, filename_output

def clean_text(text):
    # Remove emails
    text = re.sub(r'\S+@\S+', '', text)

    # Remove phone numbers (various formats)
    text = re.sub(r'\(?\+?[0-9]*\)?-?[0-9]+-?[0-9]+-?[0-9]+', '', text)

    # Remove characters that repeat more than 10 times
    text = re.sub(r'(.)\1{10,}', '', text)

    # Remove multiple whitespaces
    text = re.sub(r'\s+', ' ', text).strip()

    print(text)
    return text

llm_model = "gpt-4-1106-preview"

def read_text_from_file(file_path):
    text = ""
    print(file_path)
    if file_path.endswith(('.doc', '.docx')):
        document = docx2txt.process(file_path)
        text += document + "\n"
        num_pages = None
    elif file_path.endswith('.pdf'):
        reader = PdfReader(file_path)    
        pdf_reader = PyPDF2.PdfReader(file_path)
        num_pages = len(pdf_reader.pages)
        print(num_pages)
        for page in range(len(reader.pages)):
            text += reader.pages[page].extract_text()
    else:
        print("Error: File format not supported")
    return text, num_pages

def read_pdf(file_path, max_tokens=20000):
    text, num_pages = read_text_from_file(file_path)

    relevant_text = []

    if file_path.endswith(('.doc', '.docx')):
        # Add additional processing if needed for doc/docx
        cleaned_page_text = clean_text(text)
        relevant_text.append(cleaned_page_text)
    elif file_path.endswith('.pdf'):
        if num_pages < 3:
            print("small pdf")
            for page_num in range(num_pages):
                page_text = text.lower()  # No need to extract text again
                cleaned_page_text = clean_text(page_text)
                relevant_text.append(cleaned_page_text)
        else:
            print("big pdf")
            for page_num in range(num_pages):
                page_text = text.lower()  # No need to extract text again
                cleaned_page_text = clean_text(page_text)
                relevant_text.append(cleaned_page_text)

            if not relevant_text:
                print("No relevant page found, including as much content as possible.")
                all_text = text.lower()  # No need to extract text again
                cleaned_text = clean_text(all_text)
                relevant_text.append(cleaned_text)

    return str(" ".join(relevant_text))[:max_tokens]

def extract_client_name(jd="", email=""):
    if jd!="":
        jd=read_pdf(jd)
        print(jd)
    content_text = email + jd
    # Part 2
    client_name_description = """
    Identify the client name from the job description or email text.
    Ensure to check against the following client database:
    1: Adams 12 Five Star Schools (Adams 12)
    2: BankUnited (BU)
    3: Bayview Asset Management (Bayview)
    4: Broward College (BC)
    5: Broward County Sheriff's Office (BCSO or BSO)
    6: City of Miami (COM)
    7: City of Miami Beach (COMB)
    8: City of North Miami Beach (CONMB or CNMB)
    9: City of Fort Lauderdale (CFL)
    10: Collier County Sheriff’s Office (CCSO)
    11: Florida Atlantic University (FAU)
    12: Hillsborough County Sheriff’s Office (HCSO)
    13: Jefferson County Schools (Jeffco)
    14: Miami Police Department (MPD)
    15: Miami Dade College ()
    16: Miami-Dade County (MDC)
    17: Miami-Dade County Public Schools (MDCPS)
    18: Miami-Dade Property Appraiser (MDPA)
    19: Miami Dade Water and Sewer Department (MDWSD)
    20: Moffitt Cancer Center (Moffitt)
    21: NWRDC-FSU (NWRDC FSU)
    22: Orange County Public Schools (OCPS)
    23: Palm Beach Public Schools (PBPS)
    24: Palm Beach Sheriff's Office (PBSO)
    25: South Florida Water Management District (SFWMD)
    26: Nextera Energy Resources (NEE or FPL or Florida Power & Light)
    27: Carnival Cruise Line, non-IT (CCL)
    28: Carnival Cruise Line (CCL)
    29: Florida Agency For Health Care Administration (AHCA or FAHCA)
    30: Florida Agency for Persons with Disabilities (APD)
    31: Florida Department of Agriculture and Consumer Services (DACS)
    32: Florida Department of Business and Professional Regulation (DBPR)
    33: Florida Department of Children and Families (DCF)
    34: Florida Department of Corrections (FDC)
    35: Florida Department of Economic Opportunity (DEO)
    36: Florida Department of Education (FDOE)
    37: Florida Department of Elder Affairs (DOEA)
    38: Florida Department of Environmental Protection (DEP)
    39: Flordia Department of Financial Services (DFS)
    40: Florida Department of Health (DOH)
    41: Florida Department of Highway Safety & Motor Vehicles (FLHSMV)
    42: Florida Department of Law Enforcement (FDLE or DLFE)
    43: Florida Department of Management Services (DMS)
    44: Florida Department of Revenue (DOR)
    45: Florida Department of State (DOS)
    46: Florida Department of Transportation (FDOT or DOT)
    47: Florida Department of Transportation - Turnpike Enterprise (Turnpike)
    48: Florida Fish and Wildlife Conservation Commission (FWC)
    49: Florida Gaming Control Commission (FGCC)
    50: Florida Lottery (Lottery)
    51: Florida Virtual School (FLVS)
    52: Florida State Courts System (OSCA)
    53: Broward Health (BH)
    """

    client_name_schema = ResponseSchema(name="Client Name", description=client_name_description)
    client_id_schema = ResponseSchema(name="Client ID", description="Integer representing the ID of the client")

    response_schemas_part2 = [client_name_schema, client_id_schema]
    output_parser_part2 = StructuredOutputParser.from_response_schemas(response_schemas_part2)
    format_instructions_part2 = output_parser_part2.get_format_instructions()

    # SECTION 2
    client_name_p_template = ChatPromptTemplate.from_template(
        """
        Analyze the provided content, which can either be a job description or an email text or both. Your main objective is to identify the client name and then check it against a predefined client database.

        Extract the following:
        - Client Name and Client ID: {client_name_description}

        If Client Name found, return the name of the client and ID of the client.
        If Client Name not found : return none as the name of the client and -1 as the ID of the client.
        Content:
        ```{content_text}```

        Use this format for extraction:
        ----------------------------
        {format_instructions}
        ----------------------------
        return a valid JSON
        """
    )

    client_name_prompt = client_name_p_template.format_messages(content_text=content_text, client_name_description=client_name_description, format_instructions=format_instructions_part2)

    chat = ChatOpenAI(temperature=0.0, model=llm_model, openai_api_key=api_key, request_timeout=620)
    part2_raw = chat(client_name_prompt)

    part2 = output_parser_part2.parse(part2_raw.content)

    return part2

# email = read_pdf("/home/jarial/Desktop/Office/Pinaaki/inputs/Word File.docx")
# jd = read_pdf("/home/jarial/Desktop/Office/Pinaaki/inputs/Word File.docx")
# x = extract_client_name(email=email)
# print(x)

# print(read_pdf)